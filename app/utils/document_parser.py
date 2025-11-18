"""
Document parser for multi-format resume support

Supports:
- PDF files (.pdf) - with OCR fallback for scanned documents
- Word documents (.docx, .doc)
- Text files (.txt)
- Markdown files (.md)
"""
import logging
from pathlib import Path
from typing import Union, BinaryIO, Optional
import io

logger = logging.getLogger(__name__)


class DocumentParseError(Exception):
    """Exception raised when document parsing fails"""
    pass


class DocumentParser:
    """
    Multi-format document parser for resumes

    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    - Text (.txt)
    - Markdown (.md)
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}

    def __init__(self, ocr_settings: Optional['PdfOcrSettings'] = None):
        """
        Initialize document parser

        Args:
            ocr_settings: OCR settings for PDF parsing (optional)
        """
        self._lazy_imports_done = False
        self._ocr_settings = ocr_settings
        self._ocr_parser = None

    def _lazy_import_dependencies(self):
        """Lazy import heavy dependencies only when needed"""
        if self._lazy_imports_done:
            return

        global pypdf, docx, chardet
        try:
            import pypdf
            import docx
            import chardet
            self._lazy_imports_done = True
        except ImportError as e:
            raise DocumentParseError(
                f"Required parsing library not installed: {e}. "
                "Please run: pip install pypdf python-docx chardet"
            )

    @staticmethod
    def is_supported(file_path: Union[str, Path]) -> bool:
        """
        Check if file format is supported

        Args:
            file_path: Path to file

        Returns:
            True if format is supported, False otherwise
        """
        path = Path(file_path)
        return path.suffix.lower() in DocumentParser.SUPPORTED_EXTENSIONS

    @staticmethod
    def detect_encoding(file_path: Union[str, Path]) -> str:
        """
        Detect text file encoding

        Args:
            file_path: Path to text file

        Returns:
            Detected encoding (e.g., 'utf-8', 'gb2312', 'gbk')
        """
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
                confidence = result['confidence']

                logger.debug(f"Detected encoding: {encoding} (confidence: {confidence:.2%})")

                # Default to utf-8 if confidence is too low
                if confidence < 0.7:
                    logger.warning(
                        f"Low confidence ({confidence:.2%}) for detected encoding {encoding}, "
                        "defaulting to utf-8"
                    )
                    return 'utf-8'

                return encoding or 'utf-8'
        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}, defaulting to utf-8")
            return 'utf-8'

    def _extract_pdf_text_original(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from PDF using standard text extraction (no OCR)

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content

        Raises:
            DocumentParseError: If parsing fails
        """
        self._lazy_import_dependencies()

        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue

            full_text = '\n\n'.join(text_parts)
            logger.debug(f"Text extraction: {len(reader.pages)} pages, {len(full_text)} chars")

            return full_text

        except Exception as e:
            logger.warning(f"Text extraction failed: {e}")
            return ""

    def _get_ocr_settings(self) -> 'PdfOcrSettings':
        """Get OCR settings (lazy load from env if not provided)"""
        if self._ocr_settings is None:
            try:
                from app.utils.pdf_ocr_parser import create_ocr_settings_from_env
                self._ocr_settings = create_ocr_settings_from_env()
            except Exception as e:
                logger.warning(f"Failed to load OCR settings from env: {e}")
                # Create default settings with OCR disabled
                from app.utils.pdf_ocr_parser import PdfOcrSettings
                self._ocr_settings = PdfOcrSettings(enabled=False)

        return self._ocr_settings

    def _get_ocr_parser(self) -> 'PdfOcrParser':
        """Get OCR parser (lazy initialization)"""
        if self._ocr_parser is None:
            from app.utils.pdf_ocr_parser import PdfOcrParser
            settings = self._get_ocr_settings()
            self._ocr_parser = PdfOcrParser(settings)

        return self._ocr_parser

    def parse_pdf(self, file_path: Union[str, Path], use_ocr: Optional[bool] = None) -> str:
        """
        Parse PDF file to extract text with smart OCR fallback

        This method:
        1. First tries standard text extraction (fast)
        2. If text extraction fails or returns too little text, uses OCR
        3. Can be forced to use OCR via use_ocr=True or force_ocr setting

        Args:
            file_path: Path to PDF file
            use_ocr: Force OCR usage (None = auto-detect, True = force, False = disable)

        Returns:
            Extracted text content

        Raises:
            DocumentParseError: If parsing fails
        """
        file_path = Path(file_path)

        try:
            # Get OCR settings
            ocr_settings = self._get_ocr_settings()

            # Determine if OCR should be used
            ocr_enabled = ocr_settings.enabled if use_ocr is None else use_ocr
            force_ocr = ocr_settings.force_ocr and ocr_enabled

            # Try text extraction first (unless force_ocr is True)
            text_based = ""
            if not force_ocr:
                logger.info("Attempting text-based PDF extraction")
                text_based = self._extract_pdf_text_original(file_path)

            # Decide whether to use OCR
            use_ocr_fallback = False

            if force_ocr:
                logger.info("Force OCR is enabled, using OCR")
                use_ocr_fallback = True
            elif not ocr_enabled:
                logger.debug("OCR is disabled")
                use_ocr_fallback = False
            elif len(text_based.strip()) < ocr_settings.min_text_length:
                logger.info(
                    f"Text extraction returned {len(text_based.strip())} chars "
                    f"(< {ocr_settings.min_text_length}), falling back to OCR"
                )
                use_ocr_fallback = True
            else:
                logger.info(f"Text extraction successful ({len(text_based)} chars), OCR not needed")
                use_ocr_fallback = False

            # Use OCR if needed
            if use_ocr_fallback:
                try:
                    logger.info("Starting OCR-based extraction")
                    ocr_parser = self._get_ocr_parser()
                    ocr_text = ocr_parser.extract_text(file_path)

                    if not ocr_text or len(ocr_text.strip()) < 50:
                        if text_based:
                            logger.warning("OCR returned little text, using text extraction result")
                            return text_based
                        else:
                            raise DocumentParseError(
                                "Both text extraction and OCR failed to extract meaningful content"
                            )

                    logger.info(f"OCR extraction successful: {len(ocr_text)} chars")
                    return ocr_text

                except ImportError as e:
                    logger.warning(f"OCR dependencies not available: {e}")
                    if text_based:
                        logger.info("Falling back to text extraction result")
                        return text_based
                    else:
                        raise DocumentParseError(
                            f"OCR not available and text extraction failed. {e}"
                        )
                except Exception as e:
                    logger.error(f"OCR extraction failed: {e}")
                    if text_based:
                        logger.info("Falling back to text extraction result")
                        return text_based
                    else:
                        raise DocumentParseError(f"OCR failed and no text extraction fallback: {e}")

            # Return text-based extraction
            if not text_based:
                raise DocumentParseError("No text could be extracted from PDF")

            logger.info(f"Successfully parsed PDF: {len(text_based)} chars")
            return text_based

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse PDF file: {e}")

    def parse_docx(self, file_path: Union[str, Path]) -> str:
        """
        Parse DOCX file to extract text

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content

        Raises:
            DocumentParseError: If parsing fails
        """
        self._lazy_import_dependencies()

        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            if not text_parts:
                raise DocumentParseError("No text could be extracted from DOCX")

            full_text = '\n'.join(text_parts)
            logger.info(f"Successfully parsed DOCX: {len(doc.paragraphs)} paragraphs, {len(full_text)} chars")

            return full_text

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse DOCX file: {e}")

    def parse_text(self, file_path: Union[str, Path]) -> str:
        """
        Parse text file (TXT, MD)

        Args:
            file_path: Path to text file

        Returns:
            File content

        Raises:
            DocumentParseError: If parsing fails
        """
        try:
            # Try UTF-8 first (most common)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    logger.info(f"Successfully parsed text file with UTF-8: {len(text)} chars")
                    return text
            except UnicodeDecodeError:
                # Try auto-detection for other encodings (e.g., GBK, GB2312 for Chinese)
                encoding = self.detect_encoding(file_path)
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    logger.info(f"Successfully parsed text file with {encoding}: {len(text)} chars")
                    return text

        except Exception as e:
            raise DocumentParseError(f"Failed to parse text file: {e}")

    def parse_file(self, file_path: Union[str, Path]) -> str:
        """
        Parse file based on extension

        Args:
            file_path: Path to file

        Returns:
            Extracted text content

        Raises:
            DocumentParseError: If file format is unsupported or parsing fails
        """
        path = Path(file_path)

        if not path.exists():
            raise DocumentParseError(f"File not found: {file_path}")

        if not path.is_file():
            raise DocumentParseError(f"Not a file: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise DocumentParseError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        logger.info(f"Parsing document: {path.name} ({extension})")

        # Route to appropriate parser
        if extension == '.pdf':
            return self.parse_pdf(path)
        elif extension == '.docx':
            return self.parse_docx(path)
        elif extension in {'.txt', '.md'}:
            return self.parse_text(path)
        else:
            raise DocumentParseError(f"No parser available for {extension}")

    def parse_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse file from bytes (useful for API uploads)

        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to determine format)

        Returns:
            Extracted text content

        Raises:
            DocumentParseError: If parsing fails
        """
        extension = Path(filename).suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise DocumentParseError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        logger.info(f"Parsing uploaded document: {filename} ({extension})")

        try:
            if extension == '.pdf':
                self._lazy_import_dependencies()
                from pypdf import PdfReader

                pdf_file = io.BytesIO(file_bytes)
                reader = PdfReader(pdf_file)
                text_parts = []

                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)

                if not text_parts:
                    raise DocumentParseError("No text could be extracted from PDF")

                return '\n\n'.join(text_parts)

            elif extension == '.docx':
                self._lazy_import_dependencies()
                from docx import Document

                docx_file = io.BytesIO(file_bytes)
                doc = Document(docx_file)
                text_parts = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                for table in doc.tables:
                    for row in table.rows:
                        row_text = '\t'.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text_parts.append(row_text)

                if not text_parts:
                    raise DocumentParseError("No text could be extracted from DOCX")

                return '\n'.join(text_parts)

            elif extension in {'.txt', '.md'}:
                # Try UTF-8 first
                try:
                    return file_bytes.decode('utf-8')
                except UnicodeDecodeError:
                    # Try auto-detection
                    import chardet
                    result = chardet.detect(file_bytes)
                    encoding = result['encoding'] or 'utf-8'
                    return file_bytes.decode(encoding, errors='replace')

            else:
                raise DocumentParseError(f"No parser available for {extension}")

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse uploaded file: {e}")


# Singleton instance
document_parser = DocumentParser()


# Convenience functions
def parse_resume(file_path: Union[str, Path]) -> str:
    """
    Parse resume file (convenience function)

    Args:
        file_path: Path to resume file

    Returns:
        Extracted text content

    Raises:
        DocumentParseError: If parsing fails
    """
    return document_parser.parse_file(file_path)


def parse_resume_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Parse resume from bytes (convenience function for API uploads)

    Args:
        file_bytes: File content as bytes
        filename: Original filename

    Returns:
        Extracted text content

    Raises:
        DocumentParseError: If parsing fails
    """
    return document_parser.parse_bytes(file_bytes, filename)


def is_supported_format(filename: str) -> bool:
    """
    Check if filename has supported extension

    Args:
        filename: Filename to check

    Returns:
        True if supported, False otherwise
    """
    return DocumentParser.is_supported(filename)
